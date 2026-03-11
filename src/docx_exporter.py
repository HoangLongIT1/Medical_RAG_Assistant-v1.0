"""
docx_exporter.py – Xuất báo cáo Markdown thành file DOCX
Sử dụng python-docx — Hỗ trợ đầy đủ tiếng Việt (Unicode).
"""

import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def _strip_emoji(text: str) -> str:
    """Loại bỏ emoji và ký tự ngoài BMP."""
    text = re.sub(r'[\U00010000-\U0010FFFF]', '', text)
    text = text.replace('\ufe0f', '')
    bmp_symbols = ['\u2695', '\u26a0', '\u2705', '\u274c', '\u2b07',
                   '\u23f3', '\u23f9', '\u2b50']
    for s in bmp_symbols:
        text = text.replace(s, '')
    return text


def markdown_to_docx_bytes(md_text: str, title: str = "Báo Cáo Y Tế") -> bytes:
    """
    Chuyển Markdown text thành mảng bytes DOCX.
    Hỗ trợ: headings, bold, italic, bullet list, horizontal rules.
    """
    md_text = _strip_emoji(md_text)
    doc = Document()

    # ── Cấu hình style mặc định ──────────────────────────
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    # Đặt font cho tất cả các heading
    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Arial'
        heading_style.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    # ── Cấu hình trang ──────────────────────────────────
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    lines = md_text.split('\n')
    current_paragraph_lines = []

    def flush_paragraph():
        """Ghi đoạn văn bản tích lũy."""
        nonlocal current_paragraph_lines
        if current_paragraph_lines:
            text = ' '.join(current_paragraph_lines).strip()
            if text:
                _add_formatted_paragraph(doc, text, style='Normal')
            current_paragraph_lines = []

    for line in lines:
        stripped = line.strip()

        # Dòng trống
        if not stripped:
            flush_paragraph()
            continue

        # Horizontal rule
        if stripped == '---':
            flush_paragraph()
            _add_horizontal_rule(doc)
            continue

        # Heading 1
        if stripped.startswith('# ') and not stripped.startswith('## '):
            flush_paragraph()
            heading_text = stripped[2:].replace('**', '').strip()
            heading_text = re.sub(r'[^\w\s\-–—:()&,.]', '', heading_text).strip()
            p = doc.add_heading(heading_text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Thêm đường kẻ dưới heading 1
            _add_horizontal_rule(doc, color='0284c7', thickness=1.5)
            continue

        # Heading 3 (xử lý trước Heading 2)
        if stripped.startswith('### '):
            flush_paragraph()
            h3_text = stripped[4:].replace('**', '').strip()
            doc.add_heading(h3_text, level=3)
            continue

        # Heading 2
        if stripped.startswith('## '):
            flush_paragraph()
            h2_text = stripped[3:].replace('**', '').strip()
            p = doc.add_heading(h2_text, level=2)
            continue

        # Bullet list
        if stripped.startswith('- ') or stripped.startswith('* '):
            flush_paragraph()
            item_text = stripped[2:].strip()
            _add_formatted_paragraph(doc, item_text, style='List Bullet')
            continue

        # Regular text — tích lũy
        line_clean = stripped
        current_paragraph_lines.append(line_clean)

    flush_paragraph()

    # ── Xuất ra bytes ─────────────────────────────────────
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_formatted_paragraph(doc: Document, text: str, style: str = 'Normal'):
    """
    Thêm paragraph với hỗ trợ markdown bold (**text**) và italic (*text*).
    """
    p = doc.add_paragraph(style=style)

    # Parse bold và italic
    # Pattern: **bold** hoặc *italic*
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            # Italic
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            # Normal text
            # Xử lý thêm __bold__ nếu có
            sub_parts = re.split(r'(__.*?__)', part)
            for sub in sub_parts:
                if sub.startswith('__') and sub.endswith('__'):
                    run = p.add_run(sub[2:-2])
                    run.bold = True
                else:
                    p.add_run(sub)


def _add_horizontal_rule(doc: Document, color: str = 'cbd5e1', thickness: float = 0.5):
    """Thêm đường kẻ ngang vào document."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Dùng border bottom cho paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): str(int(thickness * 8)),
        qn('w:space'): '1',
        qn('w:color'): color,
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
